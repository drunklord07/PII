import os
import re
import shutil
from multiprocessing import Pool, cpu_count, Manager
from docx import Document
from docx.shared import RGBColor
import xlsxwriter
from tqdm import tqdm

# === CONFIG ===
CHUNK_SIZE = 2000
INPUT_FILE = "input.txt"
EMAIL_REGEX = r"(?<![\w@.-])([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})(?![\w@.-])"
OUTPUT_DOCX = "output.docx"
OUTPUT_XLSX = "output.xlsx"
TEMP_DIR = "temp_parts"

# === SPLIT FILE INTO CHUNKS ===
def split_file():
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)
    os.makedirs(TEMP_DIR)

    with open(INPUT_FILE, 'r', encoding='utf-8', errors='ignore') as infile:
        lines = []
        part = 0
        total = sum(1 for _ in open(INPUT_FILE, 'r', encoding='utf-8', errors='ignore'))
        with tqdm(total=total, desc="Splitting file", unit="lines") as pbar:
            for i, line in enumerate(infile, 1):
                lines.append(line)
                if i % CHUNK_SIZE == 0:
                    with open(f"{TEMP_DIR}/chunk_{part}.txt", 'w', encoding='utf-8') as out:
                        out.writelines(lines)
                    lines = []
                    part += 1
                pbar.update(1)
            if lines:
                with open(f"{TEMP_DIR}/chunk_{part}.txt", 'w', encoding='utf-8') as out:
                    out.writelines(lines)

# === PROCESS ONE CHUNK (WORKER) ===
def process_chunk(args):
    chunk_path, result_list = args
    chunk_id = os.path.splitext(os.path.basename(chunk_path))[0].split('_')[-1]
    pattern = re.compile(EMAIL_REGEX)

    doc = Document()
    matches_data = []
    total_lines = 0
    matched_lines = 0

    with open(chunk_path, 'r', encoding='utf-8', errors='ignore') as f:
        for line in f:
            total_lines += 1
            line = line.rstrip('\n')
            matches = list(pattern.finditer(line))
            if not matches:
                continue
            matched_lines += 1

            # Word
            para = doc.add_paragraph()
            last_idx = 0
            for match in matches:
                start, end = match.span(1)
                if start > last_idx:
                    para.add_run(line[last_idx:start])
                red_run = para.add_run(line[start:end])
                red_run.font.color.rgb = RGBColor(255, 0, 0)
                last_idx = end
            if last_idx < len(line):
                para.add_run(line[last_idx:])

            # Store for Excel later
            for match in matches:
                matches_data.append((match.group(1), line, [(m.start(1), m.end(1)) for m in pattern.finditer(line)]))

    # Save Word part
    doc.save(f"{TEMP_DIR}/chunk_{chunk_id}.docx")
    result_list.append((matches_data, total_lines, matched_lines))

# === MERGE WORD FILES ===
def merge_word():
    merged = Document()
    part_files = sorted(f for f in os.listdir(TEMP_DIR) if f.endswith(".docx"))
    with tqdm(total=len(part_files), desc="Merging Word files") as pbar:
        for file in part_files:
            doc = Document(f"{TEMP_DIR}/{file}")
            for para in doc.paragraphs:
                merged.add_paragraph(para.text)
            pbar.update(1)
    merged.save(OUTPUT_DOCX)

# === FINAL EXCEL WRITING WITH COLOR ===
def write_excel(all_matches):
    workbook = xlsxwriter.Workbook(OUTPUT_XLSX)
    worksheet = workbook.add_worksheet()
    red_format = workbook.add_format({'font_color': 'red'})

    row = 0
    for match_text, line_text, match_spans in all_matches:
        worksheet.write(row, 0, match_text, red_format)

        # Build rich string for full line
        segments = []
        last_idx = 0
        for start, end in match_spans:
            if start > last_idx:
                segments.append(line_text[last_idx:start])
            segments.append(red_format)
            segments.append(line_text[start:end])
            last_idx = end
        if last_idx < len(line_text):
            segments.append(line_text[last_idx:])
        worksheet.write_rich_string(row, 1, *segments)
        row += 1

    workbook.close()

# === MAIN ===
if __name__ == "__main__":
    print("=== EMAIL EXTRACTOR (Parallel + Colored Output) ===")
    print("Regex used:", EMAIL_REGEX)

    split_file()

    print("\nProcessing chunks in parallel...")
    manager = Manager()
    results = manager.list()

    chunk_files = [f"{TEMP_DIR}/{f}" for f in os.listdir(TEMP_DIR) if f.endswith(".txt")]
    args = [(cf, results) for cf in chunk_files]

    with tqdm(total=len(chunk_files), desc="Chunk Processing") as pbar:
        with Pool(min(cpu_count(), len(chunk_files))) as pool:
            for _ in pool.imap_unordered(process_chunk, args):
                pbar.update(1)

    # Collect data
    all_matches = []
    total_lines_scanned = 0
    total_lines_matched = 0
    for matches_data, total, matched in results:
        all_matches.extend(matches_data)
        total_lines_scanned += total
        total_lines_matched += matched

    print(f"\nScanned lines     : {total_lines_scanned}")
    print(f"Matched lines     : {total_lines_matched}")
    print(f"Total matches     : {len(all_matches)}")

    merge_word()
    write_excel(all_matches)

    print("\nCleaning up...")
    shutil.rmtree(TEMP_DIR)

    print(f"\nDONE ✅")
    print(f"→ Word Output : {OUTPUT_DOCX}")
    print(f"→ Excel Output: {OUTPUT_XLSX}")
    print("\nIn Excel:")
    print("- Column A: Matched emails (red)")
    print("- Column B: Full line with each email in red")
    print("\nIn Word:")
    print("- Each matching line as paragraph, emails highlighted in red.")
